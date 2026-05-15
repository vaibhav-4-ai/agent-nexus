"""
Document Parser — PDF and DOCX extraction.
"""

from __future__ import annotations

from dataclasses import dataclass, field
from typing import Any

from src.infra.logging import get_logger

logger = get_logger("perception.document")


@dataclass
class ExtractedDocument:
    """Extracted document content."""
    text: str
    pages: int = 0
    tables: list[list[list[str]]] = field(default_factory=list)
    images: list[bytes] = field(default_factory=list)
    metadata: dict[str, Any] = field(default_factory=dict)


class DocumentParser:
    """Parse PDF and DOCX documents into structured text."""

    async def extract_text(self, file_data: bytes, file_type: str) -> ExtractedDocument:
        """Extract text from a document."""
        file_type = file_type.lower().strip(".")
        match file_type:
            case "pdf":
                return self._parse_pdf(file_data)
            case "docx":
                return self._parse_docx(file_data)
            case "txt" | "md" | "csv":
                text = file_data.decode("utf-8", errors="replace")
                return ExtractedDocument(text=text, pages=1)
            case _:
                return ExtractedDocument(text=f"[Unsupported format: {file_type}]")

    def _parse_pdf(self, data: bytes) -> ExtractedDocument:
        """Parse PDF using PyMuPDF."""
        try:
            import fitz  # pymupdf
            doc = fitz.open(stream=data, filetype="pdf")
            pages_text = []
            tables: list[list[list[str]]] = []

            for page in doc:
                pages_text.append(page.get_text())

            full_text = "\n\n--- Page Break ---\n\n".join(pages_text)
            return ExtractedDocument(
                text=full_text,
                pages=len(doc),
                tables=tables,
                metadata={"format": "pdf", "page_count": len(doc)},
            )
        except Exception as e:
            logger.error("pdf_parse_failed", error=str(e))
            return ExtractedDocument(text=f"[PDF parse error: {e}]")

    def _parse_docx(self, data: bytes) -> ExtractedDocument:
        """Parse DOCX using python-docx."""
        try:
            import io
            from docx import Document

            doc = Document(io.BytesIO(data))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            full_text = "\n\n".join(paragraphs)

            # Extract tables
            tables = []
            for table in doc.tables:
                rows = []
                for row in table.rows:
                    rows.append([cell.text for cell in row.cells])
                tables.append(rows)

            return ExtractedDocument(
                text=full_text,
                pages=1,
                tables=tables,
                metadata={"format": "docx", "paragraphs": len(paragraphs), "tables": len(tables)},
            )
        except Exception as e:
            logger.error("docx_parse_failed", error=str(e))
            return ExtractedDocument(text=f"[DOCX parse error: {e}]")

    async def extract_tables(self, file_data: bytes, file_type: str) -> list[list[list[str]]]:
        """Extract tables from a document."""
        result = await self.extract_text(file_data, file_type)
        return result.tables
